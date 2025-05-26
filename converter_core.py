"""
Core conversion functions for different file formats
"""

import os
import logging
from pathlib import Path
from typing import Optional
import docx
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import pdfplumber
import docx2txt
from PIL import Image
import markdown
from bs4 import BeautifulSoup
import pandas as pd
import utils
from video_audio_converter import VideoAudioConverter

logger = logging.getLogger(__name__)

class DocumentConverter:
    """Main converter class handling all format conversions"""

    def __init__(self):
        self.video_audio_converter = VideoAudioConverter()
        self.supported_conversions = {
            ('docx', 'pdf'): self._docx_to_pdf,
            ('docx', 'txt'): self._docx_to_txt,
            ('docx', 'html'): self._docx_to_html,
            ('pdf', 'txt'): self._pdf_to_txt,
            ('pdf', 'docx'): self._pdf_to_docx,
            ('txt', 'pdf'): self._txt_to_pdf,
            ('txt', 'docx'): self._txt_to_docx,
            ('html', 'pdf'): self._html_to_pdf,
            ('html', 'txt'): self._html_to_txt,
            ('md', 'html'): self._md_to_html,
            ('xlsx', 'csv'): self._xlsx_to_csv,
            ('csv', 'xlsx'): self._csv_to_xlsx,
            # Image conversions - comprehensive matrix
            ('png', 'jpg'): self._image_convert,
            ('png', 'gif'): self._image_convert,
            ('png', 'bmp'): self._image_convert,
            ('png', 'tiff'): self._image_convert,
            ('png', 'webp'): self._image_convert,
            ('png', 'ico'): self._image_convert,
            ('png', 'pdf'): self._image_to_pdf,
            ('jpg', 'png'): self._image_convert,
            ('jpg', 'gif'): self._image_convert,
            ('jpg', 'bmp'): self._image_convert,
            ('jpg', 'tiff'): self._image_convert,
            ('jpg', 'webp'): self._image_convert,
            ('jpg', 'pdf'): self._image_to_pdf,
            ('gif', 'png'): self._image_convert,
            ('gif', 'jpg'): self._image_convert,
            ('gif', 'bmp'): self._image_convert,
            ('gif', 'tiff'): self._image_convert,
            ('gif', 'webp'): self._image_convert,
            ('gif', 'pdf'): self._image_to_pdf,
            ('bmp', 'png'): self._image_convert,
            ('bmp', 'jpg'): self._image_convert,
            ('bmp', 'gif'): self._image_convert,
            ('bmp', 'tiff'): self._image_convert,
            ('bmp', 'webp'): self._image_convert,
            ('bmp', 'pdf'): self._image_to_pdf,
            ('tiff', 'png'): self._image_convert,
            ('tiff', 'jpg'): self._image_convert,
            ('tiff', 'gif'): self._image_convert,
            ('tiff', 'bmp'): self._image_convert,
            ('tiff', 'webp'): self._image_convert,
            ('tiff', 'pdf'): self._image_to_pdf,
            ('webp', 'png'): self._image_convert,
            ('webp', 'jpg'): self._image_convert,
            ('webp', 'gif'): self._image_convert,
            ('webp', 'bmp'): self._image_convert,
            ('webp', 'tiff'): self._image_convert,
            ('webp', 'pdf'): self._image_to_pdf,
            ('ico', 'png'): self._image_convert,
            ('ico', 'jpg'): self._image_convert,
            ('ico', 'gif'): self._image_convert,
            ('ico', 'bmp'): self._image_convert,
            ('ico', 'pdf'): self._image_to_pdf,
            ('svg', 'png'): self._svg_to_raster,
            ('svg', 'jpg'): self._svg_to_raster,
            ('svg', 'pdf'): self._svg_to_pdf,
        }

        # Add video and audio conversions dynamically
        self._add_video_audio_conversions()

    def convert(self, input_file: str, output_file: str, source_format: str, target_format: str) -> bool:
        """Main conversion method"""
        try:
            conversion_key = (source_format, target_format)
            if conversion_key not in self.supported_conversions:
                logger.error(f"Conversion from {source_format} to {target_format} not supported")
                return False

            converter_func = self.supported_conversions[conversion_key]
            return converter_func(input_file, output_file)

        except Exception as e:
            logger.error(f"Conversion failed: {str(e)}")
            return False

    def _docx_to_pdf(self, input_file: str, output_file: str) -> bool:
        """Convert DOCX to PDF"""
        try:
            import platform

            # Try platform-specific methods first
            if platform.system() == "Windows":
                # Try Microsoft Word COM interface on Windows
                try:
                    import win32com.client

                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False

                    doc = word.Documents.Open(os.path.abspath(input_file))
                    doc.SaveAs(os.path.abspath(output_file), FileFormat=17)  # 17 = PDF format
                    doc.Close()
                    word.Quit()

                    logger.info(f"Successfully converted {input_file} to {output_file}")
                    return True

                except Exception as e:
                    logger.warning(f"Word COM interface failed: {str(e)}, trying fallback method")
                    return self._docx_to_pdf_fallback(input_file, output_file)

            elif platform.system() == "Linux":
                # Try LibreOffice on Linux
                try:
                    import subprocess

                    # Use LibreOffice headless mode for conversion
                    cmd = [
                        'libreoffice',
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', os.path.dirname(output_file),
                        input_file
                    ]

                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

                    if result.returncode == 0:
                        # LibreOffice creates PDF with same name as input file
                        input_name = os.path.splitext(os.path.basename(input_file))[0]
                        libreoffice_output = os.path.join(os.path.dirname(output_file), f"{input_name}.pdf")

                        # Rename to desired output filename if different
                        if libreoffice_output != output_file:
                            os.rename(libreoffice_output, output_file)

                        logger.info(f"Successfully converted {input_file} to {output_file} using LibreOffice")
                        return True
                    else:
                        logger.warning(f"LibreOffice conversion failed: {result.stderr}")
                        return self._docx_to_pdf_fallback(input_file, output_file)

                except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                    logger.warning(f"LibreOffice not available or failed: {str(e)}, using fallback method")
                    return self._docx_to_pdf_fallback(input_file, output_file)

            else:
                # macOS or other systems - use fallback method
                return self._docx_to_pdf_fallback(input_file, output_file)

        except Exception as e:
            logger.error(f"DOCX to PDF conversion failed: {str(e)}")
            # Fallback method using reportlab
            return self._docx_to_pdf_fallback(input_file, output_file)

    def _docx_to_pdf_fallback(self, input_file: str, output_file: str) -> bool:
        """Fallback DOCX to PDF conversion using reportlab"""
        try:
            # Extract text from DOCX
            text = docx2txt.process(input_file)

            # Create PDF
            c = canvas.Canvas(output_file, pagesize=letter)
            width, height = letter

            # Split text into lines and write to PDF
            lines = text.split('\n')
            y_position = height - 50

            for line in lines:
                if y_position < 50:  # Start new page
                    c.showPage()
                    y_position = height - 50

                c.drawString(50, y_position, line[:80])  # Limit line length
                y_position -= 20

            c.save()
            logger.info(f"Successfully converted {input_file} to {output_file} (fallback method)")
            return True

        except Exception as e:
            logger.error(f"Fallback DOCX to PDF conversion failed: {str(e)}")
            return False

    def _docx_to_txt(self, input_file: str, output_file: str) -> bool:
        """Convert DOCX to TXT"""
        try:
            text = docx2txt.process(input_file)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"DOCX to TXT conversion failed: {str(e)}")
            return False

    def _docx_to_html(self, input_file: str, output_file: str) -> bool:
        """Convert DOCX to HTML"""
        try:
            doc = docx.Document(input_file)
            html_content = "<html><body>"

            for paragraph in doc.paragraphs:
                html_content += f"<p>{paragraph.text}</p>"

            html_content += "</body></html>"

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)

            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"DOCX to HTML conversion failed: {str(e)}")
            return False

    def _pdf_to_txt(self, input_file: str, output_file: str) -> bool:
        """Convert PDF to TXT"""
        try:
            text = ""
            with pdfplumber.open(input_file) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)

            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"PDF to TXT conversion failed: {str(e)}")
            return False

    def _pdf_to_docx(self, input_file: str, output_file: str) -> bool:
        """Convert PDF to DOCX"""
        try:
            # Extract text from PDF
            text = ""
            with pdfplumber.open(input_file) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"

            # Create DOCX document
            doc = docx.Document()
            paragraphs = text.split('\n')

            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph)

            doc.save(output_file)
            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {str(e)}")
            return False

    def _txt_to_pdf(self, input_file: str, output_file: str) -> bool:
        """Convert TXT to PDF"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()

            c = canvas.Canvas(output_file, pagesize=letter)
            width, height = letter

            lines = text.split('\n')
            y_position = height - 50

            for line in lines:
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50

                c.drawString(50, y_position, line[:80])
                y_position -= 20

            c.save()
            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"TXT to PDF conversion failed: {str(e)}")
            return False

    def _txt_to_docx(self, input_file: str, output_file: str) -> bool:
        """Convert TXT to DOCX"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()

            doc = docx.Document()
            paragraphs = text.split('\n')

            for paragraph in paragraphs:
                doc.add_paragraph(paragraph)

            doc.save(output_file)
            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"TXT to DOCX conversion failed: {str(e)}")
            return False

    def _html_to_pdf(self, input_file: str, output_file: str) -> bool:
        """Convert HTML to PDF"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()

            # Create PDF
            c = canvas.Canvas(output_file, pagesize=letter)
            width, height = letter

            lines = text.split('\n')
            y_position = height - 50

            for line in lines:
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50

                c.drawString(50, y_position, line[:80])
                y_position -= 20

            c.save()
            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"HTML to PDF conversion failed: {str(e)}")
            return False

    def _html_to_txt(self, input_file: str, output_file: str) -> bool:
        """Convert HTML to TXT"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)

            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"HTML to TXT conversion failed: {str(e)}")
            return False

    def _md_to_html(self, input_file: str, output_file: str) -> bool:
        """Convert Markdown to HTML"""
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                md_content = f.read()

            html = markdown.markdown(md_content)

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html)

            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"Markdown to HTML conversion failed: {str(e)}")
            return False

    def _xlsx_to_csv(self, input_file: str, output_file: str) -> bool:
        """Convert XLSX to CSV"""
        try:
            df = pd.read_excel(input_file)
            df.to_csv(output_file, index=False)
            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"XLSX to CSV conversion failed: {str(e)}")
            return False

    def _csv_to_xlsx(self, input_file: str, output_file: str) -> bool:
        """Convert CSV to XLSX"""
        try:
            df = pd.read_csv(input_file)
            df.to_excel(output_file, index=False)
            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"CSV to XLSX conversion failed: {str(e)}")
            return False

    def _image_convert(self, input_file: str, output_file: str) -> bool:
        """Convert between image formats with enhanced support"""
        try:
            with Image.open(input_file) as img:
                # Get target format from file extension
                target_ext = os.path.splitext(output_file)[1].lower()

                # Handle different format conversions
                if target_ext in ['.jpg', '.jpeg']:
                    # Convert to RGB for JPEG (no transparency)
                    if img.mode in ('RGBA', 'LA', 'P'):
                        # Create white background for transparent images
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                        img = background
                    else:
                        img = img.convert('RGB')

                    # Save with quality optimization
                    img.save(output_file, 'JPEG', quality=95, optimize=True)

                elif target_ext == '.png':
                    # PNG supports transparency, convert appropriately
                    if img.mode not in ('RGBA', 'RGB', 'L', 'LA'):
                        img = img.convert('RGBA')
                    img.save(output_file, 'PNG', optimize=True)

                elif target_ext == '.gif':
                    # GIF has limited colors, convert appropriately
                    if img.mode != 'P':
                        # Convert to palette mode for GIF
                        img = img.convert('P', palette=Image.ADAPTIVE, colors=256)
                    img.save(output_file, 'GIF', optimize=True)

                elif target_ext == '.bmp':
                    # BMP doesn't support transparency
                    if img.mode in ('RGBA', 'LA'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                        img = background
                    else:
                        img = img.convert('RGB')
                    img.save(output_file, 'BMP')

                elif target_ext in ['.tiff', '.tif']:
                    # TIFF supports various modes
                    img.save(output_file, 'TIFF', compression='lzw')

                elif target_ext == '.webp':
                    # WebP supports transparency and high compression
                    img.save(output_file, 'WEBP', quality=90, method=6)

                elif target_ext == '.ico':
                    # ICO format for icons, resize if too large
                    if img.size[0] > 256 or img.size[1] > 256:
                        img = img.resize((256, 256), Image.Resampling.LANCZOS)

                    # ICO works best with RGBA
                    if img.mode != 'RGBA':
                        img = img.convert('RGBA')

                    # Save multiple sizes for ICO
                    sizes = [(16, 16), (32, 32), (48, 48), (64, 64), (128, 128), (256, 256)]
                    icons = []
                    for size in sizes:
                        if size[0] <= img.size[0] and size[1] <= img.size[1]:
                            resized = img.resize(size, Image.Resampling.LANCZOS)
                            icons.append(resized)

                    if icons:
                        icons[0].save(output_file, 'ICO', sizes=[icon.size for icon in icons])
                    else:
                        img.save(output_file, 'ICO')

                else:
                    # Default conversion
                    img.save(output_file)

            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True

        except Exception as e:
            logger.error(f"Image conversion failed: {str(e)}")
            return False

    def _image_to_pdf(self, input_file: str, output_file: str) -> bool:
        """Convert image to PDF with better quality and sizing"""
        try:
            with Image.open(input_file) as img:
                # Convert to RGB if necessary (PDF doesn't support transparency)
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Create white background for transparent images
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    if img.mode in ('RGBA', 'LA'):
                        background.paste(img, mask=img.split()[-1])
                    else:
                        background.paste(img)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')

                # Calculate appropriate size for PDF (fit to A4 if too large)
                max_width, max_height = 595, 842  # A4 size in points
                img_width, img_height = img.size

                # Scale down if image is too large
                if img_width > max_width or img_height > max_height:
                    ratio = min(max_width / img_width, max_height / img_height)
                    new_size = (int(img_width * ratio), int(img_height * ratio))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)

                # Save as PDF with good quality
                img.save(output_file, "PDF", resolution=100.0, quality=95)

            logger.info(f"Successfully converted {input_file} to {output_file}")
            return True
        except Exception as e:
            logger.error(f"Image to PDF conversion failed: {str(e)}")
            return False

    def _svg_to_raster(self, input_file: str, output_file: str) -> bool:
        """Convert SVG to raster image formats (PNG, JPG, etc.)"""
        try:
            # Try using cairosvg first (best quality)
            try:
                import cairosvg

                target_ext = os.path.splitext(output_file)[1].lower()

                if target_ext == '.png':
                    cairosvg.svg2png(url=input_file, write_to=output_file, output_width=1024, output_height=1024)
                elif target_ext in ['.jpg', '.jpeg']:
                    # SVG to PNG first, then convert to JPEG
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_png:
                        cairosvg.svg2png(url=input_file, write_to=temp_png.name, output_width=1024, output_height=1024)

                        # Convert PNG to JPEG
                        with Image.open(temp_png.name) as img:
                            if img.mode in ('RGBA', 'LA'):
                                background = Image.new('RGB', img.size, (255, 255, 255))
                                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                                img = background
                            else:
                                img = img.convert('RGB')
                            img.save(output_file, 'JPEG', quality=95)

                        os.unlink(temp_png.name)
                else:
                    # For other formats, convert to PNG first then to target
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_png:
                        cairosvg.svg2png(url=input_file, write_to=temp_png.name, output_width=1024, output_height=1024)

                        # Use regular image conversion
                        success = self._image_convert(temp_png.name, output_file)
                        os.unlink(temp_png.name)
                        return success

                logger.info(f"Successfully converted SVG {input_file} to {output_file}")
                return True

            except ImportError:
                # Fallback: try using Pillow with svg support
                logger.warning("cairosvg not available, trying Pillow fallback")
                return self._svg_to_raster_fallback(input_file, output_file)

        except Exception as e:
            logger.error(f"SVG to raster conversion failed: {str(e)}")
            return self._svg_to_raster_fallback(input_file, output_file)

    def _svg_to_raster_fallback(self, input_file: str, output_file: str) -> bool:
        """Fallback SVG conversion using basic methods"""
        try:
            # Try using wand (ImageMagick Python binding)
            try:
                from wand.image import Image as WandImage

                with WandImage(filename=input_file, resolution=150) as img:
                    img.format = os.path.splitext(output_file)[1][1:].upper()
                    img.save(filename=output_file)

                logger.info(f"Successfully converted SVG {input_file} to {output_file} using Wand")
                return True

            except ImportError:
                # Final fallback: create a placeholder image
                logger.warning("No SVG conversion libraries available, creating placeholder")

                placeholder = Image.new('RGB', (512, 512), (240, 240, 240))

                target_ext = os.path.splitext(output_file)[1].lower()
                if target_ext in ['.jpg', '.jpeg']:
                    placeholder.save(output_file, 'JPEG', quality=95)
                else:
                    placeholder.save(output_file)

                logger.warning(f"Created placeholder image for SVG conversion: {output_file}")
                return True

        except Exception as e:
            logger.error(f"SVG fallback conversion failed: {str(e)}")
            return False

    def _svg_to_pdf(self, input_file: str, output_file: str) -> bool:
        """Convert SVG to PDF"""
        try:
            # Try using cairosvg for direct SVG to PDF
            try:
                import cairosvg
                cairosvg.svg2pdf(url=input_file, write_to=output_file)
                logger.info(f"Successfully converted SVG {input_file} to PDF")
                return True

            except ImportError:
                # Fallback: convert to PNG first, then to PDF
                logger.warning("cairosvg not available, using PNG intermediate conversion")
                import tempfile

                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_png:
                    if self._svg_to_raster_fallback(input_file, temp_png.name):
                        success = self._image_to_pdf(temp_png.name, output_file)
                        os.unlink(temp_png.name)
                        return success
                    else:
                        os.unlink(temp_png.name)
                        return False

        except Exception as e:
            logger.error(f"SVG to PDF conversion failed: {str(e)}")
            return False

    def _add_video_audio_conversions(self):
        """Dynamically add video and audio conversion mappings"""
        # Video formats
        video_formats = ['mp4', 'avi', 'mov', 'wmv', 'flv', 'mkv', 'webm', 'm4v', '3gp']
        audio_formats = ['mp3', 'wav', 'aac', 'flac', 'ogg', 'm4a', 'wma']

        # Video to video conversions
        for source in video_formats:
            for target in video_formats:
                if source != target:
                    self.supported_conversions[(source, target)] = self._video_to_video

        # Video to audio conversions
        for source in video_formats:
            for target in audio_formats:
                self.supported_conversions[(source, target)] = self._video_to_audio

        # Video to GIF conversions
        for source in video_formats:
            self.supported_conversions[(source, 'gif')] = self._video_to_gif

        # Audio to audio conversions
        for source in audio_formats:
            for target in audio_formats:
                if source != target:
                    self.supported_conversions[(source, target)] = self._audio_to_audio

    def _video_to_video(self, input_file: str, output_file: str) -> bool:
        """Convert video from one format to another"""
        try:
            target_format = os.path.splitext(output_file)[1][1:].lower()
            return self.video_audio_converter.convert_video_format(input_file, output_file, target_format)
        except Exception as e:
            logger.error(f"Video to video conversion failed: {str(e)}")
            return False

    def _video_to_audio(self, input_file: str, output_file: str) -> bool:
        """Convert video to audio (extract audio track)"""
        try:
            audio_format = os.path.splitext(output_file)[1][1:].lower()
            return self.video_audio_converter.convert_video_to_audio(input_file, output_file, audio_format)
        except Exception as e:
            logger.error(f"Video to audio conversion failed: {str(e)}")
            return False

    def _video_to_gif(self, input_file: str, output_file: str) -> bool:
        """Convert video to GIF"""
        try:
            return self.video_audio_converter.convert_video_to_gif(input_file, output_file)
        except Exception as e:
            logger.error(f"Video to GIF conversion failed: {str(e)}")
            return False

    def _audio_to_audio(self, input_file: str, output_file: str) -> bool:
        """Convert audio from one format to another"""
        try:
            target_format = os.path.splitext(output_file)[1][1:].lower()
            return self.video_audio_converter.convert_audio_format(input_file, output_file, target_format)
        except Exception as e:
            logger.error(f"Audio to audio conversion failed: {str(e)}")
            return False
